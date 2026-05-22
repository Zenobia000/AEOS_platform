"""KbIngestProcessor 整合測試 — stub embedder + 真實 DB."""

from __future__ import annotations

import io
import uuid

import pytest
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.models.ingestion_job import IngestionJob
from app.db.models.knowledge_card import KnowledgeCard
from app.db.models.tenant import Tenant
from app.embeddings import StubEmbeddingClient
from app.worker.kb_ingest import (
    FileLoader,
    KbIngestProcessor,
    chunk_text,
    find_pending_ingestion_jobs,
)

# ── chunk_text ──────────────────────────────────────


def test_chunk_text_basic() -> None:
    text = "a" * 1000
    chunks = chunk_text(text, chunk_size=400, overlap=40)
    assert len(chunks) == 3  # 400 + (400-40)*2 = 1120 ≥ 1000
    assert all(len(c) <= 400 for c in chunks)


def test_chunk_text_empty() -> None:
    assert chunk_text("") == []
    assert chunk_text("   \n\t  ") == []


def test_chunk_text_smaller_than_chunk_size() -> None:
    chunks = chunk_text("hi", chunk_size=100, overlap=10)
    assert chunks == ["hi"]


def test_chunk_text_validates_params() -> None:
    with pytest.raises(ValueError):
        chunk_text("x", chunk_size=0)
    with pytest.raises(ValueError):
        chunk_text("x" * 100, chunk_size=10, overlap=10)  # overlap >= chunk
    with pytest.raises(ValueError):
        chunk_text("x", chunk_size=10, overlap=-1)


# ── Helpers ─────────────────────────────────────────


async def _make_tenant_and_job(
    session: AsyncSession,
    *,
    filename: str = "kb.txt",
    file_ref: str = "s3://bucket/kb.txt",
    slug: str = "kb",
) -> tuple[Tenant, IngestionJob]:
    t = Tenant(name="T", slug=f"{slug}-{uuid.uuid4().hex[:6]}")
    session.add(t)
    await session.flush()
    job = IngestionJob(
        tenant_id=t.id,
        source_file_ref=file_ref,
        source_filename=filename,
        status="pending",
    )
    session.add(job)
    await session.flush()
    return t, job


def _file_loader(payload: bytes) -> FileLoader:
    """Build a FileLoader callable that returns the given bytes regardless of ref."""

    async def _load(ref: str) -> bytes:
        return payload

    return _load


# ── process_one happy path ─────────────────────────


async def test_ingest_txt_creates_cards(db_session: AsyncSession) -> None:
    tenant, job = await _make_tenant_and_job(db_session, filename="faq.txt")

    payload = ("退貨期限 7 天內可申請。\n" * 80).encode("utf-8")  # ~長文字
    processor = KbIngestProcessor(
        embedder=StubEmbeddingClient(),
        file_loader=_file_loader(payload),
        chunk_size=400,
        chunk_overlap=40,
    )
    result = await processor.process_one(db_session, job)

    assert result.status == "completed"
    assert result.cards_created > 0
    assert result.error is None

    # job updated
    await db_session.refresh(job)
    assert job.status == "completed"
    assert job.cards_created == result.cards_created
    assert job.completed_at is not None

    # KC rows
    cards = (
        (
            await db_session.execute(
                select(KnowledgeCard).where(KnowledgeCard.tenant_id == tenant.id)
            )
        )
        .scalars()
        .all()
    )
    assert len(cards) == result.cards_created
    for kc in cards:
        assert kc.status == "draft"
        assert kc.card_type == "faq"
        assert kc.source_file_ref == job.source_file_ref
        assert kc.embedding is not None and len(kc.embedding) == 1024
        assert kc.embedding_model == "stub-sha256"


async def test_ingest_md(db_session: AsyncSession) -> None:
    tenant, job = await _make_tenant_and_job(db_session, filename="readme.md", slug="md")
    payload = "# 標題\n\n本店退貨政策：7 天內可申請。".encode()
    processor = KbIngestProcessor(
        embedder=StubEmbeddingClient(),
        file_loader=_file_loader(payload),
        chunk_size=200,
        chunk_overlap=20,
    )
    result = await processor.process_one(db_session, job)
    assert result.status == "completed"
    cards = (
        (
            await db_session.execute(
                select(KnowledgeCard).where(KnowledgeCard.tenant_id == tenant.id)
            )
        )
        .scalars()
        .all()
    )
    assert len(cards) >= 1
    assert "退貨" in cards[0].body_markdown


async def test_ingest_docx(db_session: AsyncSession) -> None:
    """真實 DOCX bytes 走 python-docx parser."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("第一段：營業時間週一到週五 9~18 點")
    doc.add_paragraph("第二段：退貨期限為 7 天")
    buf = io.BytesIO()
    doc.save(buf)
    payload = buf.getvalue()

    _, job = await _make_tenant_and_job(db_session, filename="hours.docx", slug="docx")
    processor = KbIngestProcessor(
        embedder=StubEmbeddingClient(),
        file_loader=_file_loader(payload),
        chunk_size=300,
        chunk_overlap=30,
    )
    result = await processor.process_one(db_session, job)
    assert result.status == "completed"
    assert result.cards_created >= 1


# ── error paths ────────────────────────────────────


async def test_ingest_unsupported_format_fails(db_session: AsyncSession) -> None:
    _, job = await _make_tenant_and_job(db_session, filename="data.csv", slug="bad-ext")
    processor = KbIngestProcessor(
        embedder=StubEmbeddingClient(),
        file_loader=_file_loader(b"a,b,c"),
    )
    result = await processor.process_one(db_session, job)
    assert result.status == "failed"
    assert result.error is not None
    assert "unsupported" in (result.error or "").lower()
    await db_session.refresh(job)
    assert job.status == "failed"
    assert job.error_message is not None


async def test_ingest_file_loader_raise_marks_failed(db_session: AsyncSession) -> None:
    _, job = await _make_tenant_and_job(db_session, filename="x.txt", slug="loader-err")

    async def bad_loader(ref: str) -> bytes:
        raise FileNotFoundError(f"missing: {ref}")

    processor = KbIngestProcessor(
        embedder=StubEmbeddingClient(),
        file_loader=bad_loader,
    )
    result = await processor.process_one(db_session, job)
    assert result.status == "failed"
    assert "missing" in (result.error or "")
    await db_session.refresh(job)
    assert job.status == "failed"


async def test_ingest_empty_file_completes_with_zero_cards(
    db_session: AsyncSession,
) -> None:
    """空白 txt → completed + cards_created=0（不算錯誤）."""
    _, job = await _make_tenant_and_job(db_session, filename="empty.txt", slug="empty")
    processor = KbIngestProcessor(
        embedder=StubEmbeddingClient(),
        file_loader=_file_loader(b"   \n  \n  "),
    )
    result = await processor.process_one(db_session, job)
    assert result.status == "completed"
    assert result.cards_created == 0


# ── pending job poller ─────────────────────────────


async def test_find_pending_ingestion_jobs(db_session: AsyncSession) -> None:
    _, j1 = await _make_tenant_and_job(db_session, slug="pend1")
    _, j2 = await _make_tenant_and_job(db_session, slug="pend2")

    pending = await find_pending_ingestion_jobs(db_session, limit=10)
    ids = {p.id for p in pending}
    assert j1.id in ids and j2.id in ids


async def test_find_pending_skips_processing_completed(db_session: AsyncSession) -> None:
    _, j_pending = await _make_tenant_and_job(db_session, slug="skip-pending")
    _, j_done = await _make_tenant_and_job(db_session, slug="skip-done")
    j_done.status = "completed"
    _, j_proc = await _make_tenant_and_job(db_session, slug="skip-proc")
    j_proc.status = "processing"
    await db_session.flush()

    pending = await find_pending_ingestion_jobs(db_session, limit=10)
    ids = {p.id for p in pending}
    assert j_pending.id in ids
    assert j_done.id not in ids
    assert j_proc.id not in ids
